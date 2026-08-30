"""Parser and chunker tests.

Fixtures are generated as real .docx/.xlsx files so the tests exercise the
actual OOXML path rather than a convenient string. Set ASSETCUES_CORPUS to a
folder of real documents to additionally run the whole-corpus checks.
"""

from __future__ import annotations

import os
from pathlib import Path

import pytest

from app.ingest.chunker import TARGET_TOKENS, chunk_markdown, count_tokens
from app.ingest.parsers import (
    UnsupportedDocument,
    _clean_filename_title,
    _derive_doc_type,
    _split_audience,
    extract_declared_audience,
    parse_document,
)

# ---------------------------------------------------------------------------
# Fixture builders
# ---------------------------------------------------------------------------


def write_docx(path: Path) -> Path:
    from docx import Document as Docx

    doc = Docx()
    doc.add_heading("User Access & Permission Management", level=1)
    doc.add_paragraph("Product & Functional Specification")

    meta = doc.add_table(rows=2, cols=3)
    meta.rows[0].cells[0].text = "Document purpose"
    meta.rows[0].cells[1].text = "Primary audience"
    meta.rows[0].cells[2].text = "Status"
    meta.rows[1].cells[0].text = "Defines the access-control model."
    meta.rows[1].cells[1].text = (
        "Product, Engineering, QA, Implementation, Support, Security and Audit"
    )
    meta.rows[1].cells[2].text = "Version 1.0"

    doc.add_heading("4. Functional requirements", level=1)
    doc.add_heading("4.4 Profiles and active access context", level=2)

    table = doc.add_table(rows=4, cols=2)
    table.rows[0].cells[0].text = "Requirement"
    table.rows[0].cells[1].text = "Confirmed current rule"
    for index, (rid, rule) in enumerate(
        [
            ("UAP-FR-037", "Create every Profile for exactly one Access Category."),
            ("UAP-FR-038", "Allow a Profile to contain Permission Groups."),
            ("UAP-FR-045", "Allow an authorized administrator to edit a Profile."),
        ],
        start=1,
    ):
        table.rows[index].cells[0].text = rid
        table.rows[index].cells[1].text = rule

    doc.save(path)
    return path


def write_xlsx(path: Path) -> Path:
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = "Test Cases"
    ws.append(["Test Case ID", "Requirement", "Scenario", "Expected Result"])
    ws.append(["UAP-TC-047", "UAP-FR-032", "Cross Legal Entity search", "Unavailable"])
    wb.save(path)
    return path


@pytest.fixture
def docx_file(tmp_path: Path) -> Path:
    return write_docx(tmp_path / "01_User_Access_Product_and_Functional_Specification.docx")


@pytest.fixture
def xlsx_file(tmp_path: Path) -> Path:
    return write_xlsx(tmp_path / "02_User_Access_Test_Cases.xlsx")


# ---------------------------------------------------------------------------
# Parsing
# ---------------------------------------------------------------------------


def test_docx_headings_become_markdown_headings(docx_file: Path) -> None:
    parsed = parse_document(docx_file)
    assert "# User Access & Permission Management" in parsed.markdown
    assert "## 4.4 Profiles and active access context" in parsed.markdown


def test_docx_tables_become_markdown_tables(docx_file: Path) -> None:
    """Requirement id and rule must stay on the same row."""
    parsed = parse_document(docx_file)
    assert (
        "| UAP-FR-045 | Allow an authorized administrator to edit a Profile. |"
        in parsed.markdown
    )


def test_declared_audience_is_extracted_from_the_template(docx_file: Path) -> None:
    parsed = parse_document(docx_file)
    assert "Product" in parsed.declared_audience
    assert "Engineering" in parsed.declared_audience
    # "Security and Audit" is the final cell entry and must split into two.
    assert "Security" in parsed.declared_audience
    assert "Audit" in parsed.declared_audience


def test_audience_split_keeps_compound_names_intact() -> None:
    """'Group and Legal Entity Administrators' is one audience, not two."""
    parts = _split_audience(
        "Group and Legal Entity Administrators, Approvers, Support and Customer Success"
    )
    assert "Group and Legal Entity Administrators" in parts
    assert "Support" in parts
    assert "Customer Success" in parts


def test_xlsx_sheets_are_parsed(xlsx_file: Path) -> None:
    parsed = parse_document(xlsx_file)
    assert "## Sheet: Test Cases" in parsed.markdown
    assert "UAP-TC-047" in parsed.markdown


def test_xlsx_title_comes_from_the_cleaned_filename(xlsx_file: Path) -> None:
    assert parse_document(xlsx_file).title == "User Access Test Cases"


def test_unsupported_type_is_rejected(tmp_path: Path) -> None:
    bad = tmp_path / "notes.rtf"
    bad.write_text("hello")
    with pytest.raises(UnsupportedDocument):
        parse_document(bad)


# ---------------------------------------------------------------------------
# Filename heuristics
# ---------------------------------------------------------------------------


@pytest.mark.parametrize(
    ("filename", "expected"),
    [
        ("01_Approval_Workflow_Product_and_Functional_Specification.docx",
         "Product & Functional Specification"),
        ("02_Approval_Workflow_Test_Cases.xlsx", "Test Cases"),
        ("03_OSM_Lean_Validation_Traceability_and_Governance_Pack.docx",
         "Validation & Governance Pack"),
        ("AssetCues_License_Management_BRD_v1.1.docx",
         "Business Requirements Document"),
        ("Reporting Period Management - User Manual.docx", "User Manual"),
        ("03_UAM_User_and_Administrator_Guide.docx", "User & Administrator Guide"),
    ],
)
def test_doc_type_is_derived_from_the_filename(filename: str, expected: str) -> None:
    """Body text says 'validation' and 'requirements' constantly, so the
    filename has to win."""
    assert _derive_doc_type(filename, "") == expected


def test_brd_is_recognised_despite_underscore_separators() -> None:
    """`\\b` does not fire between '_' and 'b'; separators are normalised."""
    assert _derive_doc_type("AssetCues_License_Management_BRD_v1.1.docx", "") == (
        "Business Requirements Document"
    )


@pytest.mark.parametrize(
    ("stem", "expected"),
    [
        ("01_OSM_Lean_Product_and_Functional_Specification (3)",
         "OSM Lean Product and Functional Specification"),
        ("AssetCues_License_Management_User_Manual_v1.1",
         "AssetCues License Management User Manual"),
        ("Field & Screen Configuration - BRD", "Field & Screen Configuration - BRD"),
    ],
)
def test_filename_titles_are_cleaned(stem: str, expected: str) -> None:
    assert _clean_filename_title(Path(f"{stem}.docx")) == expected


def test_missing_audience_field_yields_an_empty_list() -> None:
    assert extract_declared_audience("# Title\n\nSome prose.") == []


# ---------------------------------------------------------------------------
# Chunking
# ---------------------------------------------------------------------------


def test_chunks_carry_their_heading_path(docx_file: Path) -> None:
    chunks = chunk_markdown(parse_document(docx_file).markdown)
    paths = {c.heading_path for c in chunks}
    assert any("4.4 Profiles" in p for p in paths)


def test_requirement_ids_keep_their_rule_text(docx_file: Path) -> None:
    """The property that makes an answer correct."""
    chunks = chunk_markdown(parse_document(docx_file).markdown)
    hit = next(c for c in chunks if "UAP-FR-045" in c.text)
    assert "authorized administrator" in hit.text


def test_no_chunk_is_empty(docx_file: Path) -> None:
    chunks = chunk_markdown(parse_document(docx_file).markdown)
    assert chunks
    assert all(c.text.strip() for c in chunks)


def test_ordinals_are_contiguous_from_zero(docx_file: Path) -> None:
    chunks = chunk_markdown(parse_document(docx_file).markdown)
    assert [c.ordinal for c in chunks] == list(range(len(chunks)))


def test_a_large_table_is_split_on_row_boundaries_with_a_repeated_header() -> None:
    header = "| Requirement | Rule |\n|---|---|"
    rows = "\n".join(
        f"| REQ-{i:03d} | " + ("policy text " * 30) + "|" for i in range(40)
    )
    chunks = chunk_markdown(f"# Requirements\n\n{header}\n{rows}")

    assert len(chunks) > 1, "an oversized table should be split"
    for chunk in chunks:
        assert chunk.text.lstrip().startswith("| Requirement | Rule |"), (
            "every piece of a split table must repeat the header so it stays "
            "self-describing"
        )
        # A row must never be cut in half.
        for line in chunk.text.splitlines():
            if line.startswith("| REQ-"):
                assert line.rstrip().endswith("|")


def test_chunk_hash_changes_when_the_heading_changes() -> None:
    """Moving a paragraph to a new section changes its meaning."""
    a = chunk_markdown("# Alpha\n\nSame body text here.")[0]
    b = chunk_markdown("# Beta\n\nSame body text here.")[0]
    assert a.text == b.text
    assert a.sha256 != b.sha256


def test_identical_content_hashes_identically() -> None:
    """This is what makes re-ingestion skip the embedding call."""
    md = "# Alpha\n\nSame body text here."
    assert chunk_markdown(md)[0].sha256 == chunk_markdown(md)[0].sha256


def test_token_counting_is_monotonic() -> None:
    assert count_tokens("word " * 100) > count_tokens("word " * 10)


def test_empty_document_yields_no_chunks() -> None:
    assert chunk_markdown("") == []


# ---------------------------------------------------------------------------
# Whole-corpus checks (opt-in)
# ---------------------------------------------------------------------------

CORPUS = os.environ.get("ASSETCUES_CORPUS", "")
corpus_only = pytest.mark.skipif(
    not CORPUS or not Path(CORPUS).exists(),
    reason="set ASSETCUES_CORPUS to a folder of real documents",
)


@corpus_only
def test_every_real_document_parses_and_chunks() -> None:
    root = Path(CORPUS)
    files = [
        p for p in root.rglob("*")
        if p.suffix.lower() in {".docx", ".xlsx"} and not p.name.startswith("~$")
    ]
    assert files

    for path in files:
        parsed = parse_document(path)
        chunks = chunk_markdown(parsed.markdown)
        assert chunks, f"{path.name} produced no chunks"
        assert parsed.title, f"{path.name} produced no title"
        assert parsed.doc_type != "Document", f"{path.name} was not classified"
        assert all(c.text.strip() for c in chunks)
        assert max(c.token_count for c in chunks) < TARGET_TOKENS * 2
